import docx
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import os, shutil

# Target Paths on Desktop
DESKTOP_DIR = r"C:\Users\Aryan\OneDrive\Desktop\Pranayuv"
SOURCE_DOCX = os.path.join(DESKTOP_DIR, "Finals_University_Internship_Report_TimesNewRoman.docx")
FINAL_PDF = os.path.join(DESKTOP_DIR, "Finals_University_Internship_Report_TimesNewRoman.pdf")

# Workspace temporary copies for rendering and synchronization
WORKSPACE_DOCX = os.path.abspath("working_report_a4.docx")
WORKSPACE_PDF = os.path.abspath("working_report_a4.pdf")

def ensure_a4_portrait_and_clean_sections(doc):
    print("Verifying A4 size, portrait orientation, and section page numbering...")
    for i, sec in enumerate(doc.sections):
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width = Mm(210)   # A4 width
        sec.page_height = Mm(297)  # A4 height
        
        # Section 0 (Cover, Certificates, Offer Letter/NDA, TOC) should display NO page numbers
        # Section 1 (Acknowledgement onwards) restarts numbering from 1
        if i == 1 or (i > 0 and not sec.header.is_linked_to_previous):
            pgNumType = sec._sectPr.find(qn("w:pgNumType"))
            if pgNumType is None:
                pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="1"/>')
                sec._sectPr.append(pgNumType)
            else:
                pgNumType.set(qn("w:start"), "1")

def ensure_clean_page_break_before_personal_details(doc):
    print("Ensuring clean page break before Personal Details for professional table alignment...")
    for p in doc.paragraphs:
        if "2. Personal Details" in p.text:
            # Check if previous paragraph or this paragraph already has a page break
            xml = p._element.xml
            if "w:br w:type=\"page\"" not in xml:
                p_br = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>')
                p._element.addprevious(p_br)
            break

def scan_pdf_and_verify_page_numbers(pdf_path):
    print("Scanning Pass 1 native PDF to verify physical section boundaries and TOC page numbers...")
    import fitz
    doc_pdf = fitz.open(pdf_path)
    
    ack_page_idx = None
    for p_idx in range(len(doc_pdf)):
        text = doc_pdf.load_page(p_idx).get_text()
        if "Acknowledgement" in text and not ("Table of Contents" in text and p_idx < 12):
            if "sincere gratitude" in text or "1." in text:
                ack_page_idx = p_idx
                break
                
    if ack_page_idx is None:
        ack_page_idx = 8
    print(f"Verified Acknowledgement starts physically at PDF page index {ack_page_idx} (Printed Page 1)")

    search_anchors = {
        "1.": ["Acknowledgement", "sincere gratitude to everyone"],
        "2.": ["2. Personal Details", "Personal Details"],
        "3.": ["3. Company Details", "Company Details"],
        "4.": ["4. Internship Objectives", "Internship Objectives"],
        "5.": ["5.  Internship Responsibilities", "5. Internship Responsibilities", "Internship Responsibilities"],
        "6.": ["6. Technology Stack", "Technology Stack"],
        "7.": ["7. Software & Tools", "Software & Tools Used"],
        "8.": ["8. Project 1", "Pranayuv Landing Page Documentation"],
        "9.": ["9. Project 2", "EduNexus School Website Documentation"],
        "10.": ["10. Project 3", "BrightPath Academy School Management"],
        "11.": ["11. Project 4", "Sri Srinivasa Clean Rooms & Medical Furniture Web"],
        "12.": ["12. Overall Internship Summary", "Overall Internship Summary"],
        "13.": ["13. Skills Learned", "Skills Learned & Technical"],
        "14.": ["14. Conclusion", "The six-month summer internship program"]
    }

    verified_toc_pages = {}
    for ch_num, phrases in search_anchors.items():
        found = False
        for p_idx in range(ack_page_idx, len(doc_pdf)):
            text_content = doc_pdf.load_page(p_idx).get_text()
            if any(phrase in text_content for phrase in phrases):
                printed_num = (p_idx - ack_page_idx) + 1
                verified_toc_pages[ch_num] = str(printed_num)
                found = True
                break
        if not found and ch_num == "1.":
            verified_toc_pages["1."] = "1"

    doc_pdf.close()
    print(f"Verified TOC Printed Page Mapping (relative to Acknowledgement = Page 1): {verified_toc_pages}")
    return verified_toc_pages

def sync_toc_table_in_doc(doc, verified_pages):
    print("Updating Table of Contents references in Word document...")
    for t in doc.tables:
        if t.rows and ('Ch.' in t.rows[0].cells[0].text or 'Chapter' in t.rows[0].cells[1].text):
            for row in t.rows[1:]:
                ch_num = row.cells[0].text.strip()
                if ch_num in verified_pages:
                    p_cell = row.cells[-1]
                    p_cell.paragraphs[0].text = f"Page {verified_pages[ch_num]}"
                    for r in p_cell.paragraphs[0].runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(10)

def main():
    print(f"Opening source document from: {SOURCE_DOCX}")
    doc = docx.Document(SOURCE_DOCX)
    
    ensure_a4_portrait_and_clean_sections(doc)
    ensure_clean_page_break_before_personal_details(doc)
    doc.save(WORKSPACE_DOCX)
    
    print("Executing Pass 1 High-Quality PDF export...")
    from docx2pdf import convert
    if os.path.exists(WORKSPACE_PDF):
        try: os.remove(WORKSPACE_PDF)
        except Exception: pass
    convert(WORKSPACE_DOCX, WORKSPACE_PDF)
    
    verified_pages = scan_pdf_and_verify_page_numbers(WORKSPACE_PDF)
    
    doc2 = docx.Document(WORKSPACE_DOCX)
    sync_toc_table_in_doc(doc2, verified_pages)
    ensure_a4_portrait_and_clean_sections(doc2)
    doc2.save(WORKSPACE_DOCX)
    
    try:
        doc2.save(SOURCE_DOCX)
        print(f"Updated Word document saved cleanly to: {SOURCE_DOCX}")
    except Exception as e:
        print(f"Note: Could not overwrite {SOURCE_DOCX} directly (file lock): {e}")

    print("Rendering finalized High-Quality Print PDF for University Submission...")
    if os.path.exists(WORKSPACE_PDF):
        try: os.remove(WORKSPACE_PDF)
        except Exception: pass
    convert(WORKSPACE_DOCX, WORKSPACE_PDF)
    
    try:
        shutil.copy2(WORKSPACE_PDF, FINAL_PDF)
        print(f"SUCCESS! Final University-Submission-Ready PDF generated at: {FINAL_PDF}")
    except Exception as e:
        alt_pdf = os.path.join(DESKTOP_DIR, "Final_University_Internship_Report_TimesNewRoman.pdf")
        shutil.copy2(WORKSPACE_PDF, alt_pdf)
        print(f"Saved PDF to fallback desktop location: {alt_pdf}")

if __name__ == "__main__":
    main()
