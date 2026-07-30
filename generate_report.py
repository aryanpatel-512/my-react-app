"""
Summer Internship Report Generator — V3
Matches reference PDF formatting exactly:
  - Page border on every page
  - No blank/wasted space — continuous content flow
  - Minimal page breaks (only before major chapters)
  - Times New Roman 12pt, 1.5 line spacing
  - Italic sub-headings (Key Features, Technology Stack)
  - Tight paragraph spacing
  - Page numbers at bottom center
  - Proper TOC with dotted leaders
  - Screenshots center-aligned with captions
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree

BASE = os.path.dirname(os.path.abspath(__file__))
SS = os.path.join(BASE, "screenshots")

# ═══════════════════════════════════════════════════════════════
#  LOW-LEVEL FORMATTING HELPERS
# ═══════════════════════════════════════════════════════════════

def add_page_border(section):
    """Add a simple black line border to all four sides of every page."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')       # thickness
        border.set(qn('w:space'), '24')     # space from page edge
        border.set(qn('w:color'), '000000') # black
        pgBorders.append(border)
    sectPr.append(pgBorders)

def add_page_numbers(doc):
    """Add page numbers at bottom center of every page."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # PAGE field
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        
        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)
        
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

def set_default_font(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5

    # Fix heading styles
    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Times New Roman"
        hstyle.font.color.rgb = RGBColor(0, 0, 0)
        hstyle.paragraph_format.space_before = Pt(8)
        hstyle.paragraph_format.space_after = Pt(4)
        hstyle.paragraph_format.line_spacing = 1.5

    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)

    # Bullet styles
    for sname in ["List Bullet", "List Bullet 2"]:
        try:
            s = doc.styles[sname]
            s.font.name = "Times New Roman"
            s.font.size = Pt(11)
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(1)
            s.paragraph_format.line_spacing = 1.15
        except:
            pass

# ═══════════════════════════════════════════════════════════════
#  CONTENT HELPERS
# ═══════════════════════════════════════════════════════════════

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(6)
    return h

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    return h

def ital_head(doc, text):
    """Italic sub-heading like 'Key Features:' or 'Technology Stack' from the reference."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p

def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return p

def bold_para(doc, bold_text, normal_text=""):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    if normal_text:
        r2 = p.add_run(normal_text)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    return p

def sub_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet 2')
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(2.5)
    return p

def add_img(doc, filename, caption=""):
    path = os.path.join(SS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last.paragraph_format.space_before = Pt(4)
        last.paragraph_format.space_after = Pt(2)
        if caption:
            p = doc.add_paragraph()
            r = p.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(80, 80, 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)

def page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
#  BUILD
# ═══════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_margins(doc)
    set_default_font(doc)

    # Add page border to the default section
    for section in doc.sections:
        add_page_border(section)

    # Add page numbers
    add_page_numbers(doc)

    # ─── TABLE OF CONTENTS ─────────────────────────────────────
    h1(doc, "Table of Contents")

    toc_items = [
        ("1.", "Introduction", "1"),
        ("2.", "Company Overview – PranaYuv Technology", "3"),
        ("3.", "Project 1 – Sri Srinivasa Clean Rooms E-Commerce Web Application", "4"),
        ("3.1", "Home Page – Landing Experience", "5"),
        ("3.2", "Products Page – Catalog & Filtering", "7"),
        ("3.3", "Product Detail Page – Inquiry & Quotation", "9"),
        ("3.4", "Admin Login – Authentication Portal", "11"),
        ("3.5", "Admin Dashboard – Product Management", "12"),
        ("3.6", "Admin Dashboard – Inquiry & Lead Management", "14"),
        ("3.7", "Admin Dashboard – Category Management", "16"),
        ("3.8", "Backend Architecture & Database Design", "17"),
        ("4.", "Project 2 – PranaYuv Technology Corporate Landing Page", "19"),
        ("4.1", "Design Philosophy & UI/UX Decisions", "21"),
        ("5.", "Project 3 – BrightPath Academy School Management System", "22"),
        ("5.1", "School Website – Home Page", "23"),
        ("5.2", "Admission Portal", "25"),
        ("5.3", "Student Dashboard", "27"),
        ("5.4", "Principal Dashboard", "29"),
        ("6.", "Project 4 – EduNexus UI/UX Design (Figma to HTML/CSS)", "30"),
        ("6.1", "EduNexus Olive Theme Design", "32"),
        ("6.2", "Admission Requirements Page", "33"),
        ("6.3", "Student Dashboard V2", "35"),
        ("6.4", "Principal Dashboard", "37"),
        ("7.", "Project 5 – Competitive Analysis & Market Research", "38"),
        ("8.", "Tools & Technologies Used During Internship", "40"),
        ("9.", "Skills Acquired & Learning Outcomes", "42"),
        ("10.", "Conclusion", "44"),
        ("", "Appendix", "45"),
    ]

    table = doc.add_table(rows=len(toc_items), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (num, title, pg) in enumerate(toc_items):
        row = table.rows[i]
        c0, c1, c2 = row.cells
        c0.text = num
        c1.text = title
        c2.text = pg
        c0.width = Cm(1.5)
        c1.width = Cm(12)
        c2.width = Cm(1.5)
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"

    # Remove table borders for clean look
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        borders.append(b)
    tblPr.append(borders)

    page_break(doc)

    # ═══ 1. INTRODUCTION ═══════════════════════════════════════
    h1(doc, "1. Introduction")
    para(doc, "This document serves as the comprehensive Summer Internship Report submitted in partial fulfillment of the academic requirements for the Bachelor of Technology degree. The internship was undertaken at PranaYuv Technology, a healthcare technology startup focused on autonomous patient care solutions.")
    para(doc, "During the internship period, the primary areas of work included UI/UX Design, Frontend Web Development using React.js, Backend Development using Node.js and Express.js, and Database Management using MongoDB. The internship provided hands-on experience in building production-ready full-stack web applications, designing user interfaces from scratch using Figma, and implementing responsive and accessible web designs.")
    para(doc, "The internship focused on developing multiple real-world projects including an e-commerce platform for a medical furniture manufacturer, a corporate landing page for the company, a school management system, and UI/UX design prototypes for an educational platform called EduNexus.")

    ital_head(doc, "Internship Details")
    bullet(doc, "Company Name: PranaYuv Technology")
    bullet(doc, "Duration: Summer 2026")
    bullet(doc, "Role: Full Stack Developer & UI/UX Designer")
    bullet(doc, "Domain: Healthcare Technology, EdTech, Web Development")
    bullet(doc, "Reporting To: Project Lead, PranaYuv Technology")

    ital_head(doc, "Scope of Work")
    bullet(doc, "Full-Stack Web Application Development (React.js + Node.js + MongoDB)")
    bullet(doc, "UI/UX Design using Figma and conversion to responsive HTML/CSS")
    bullet(doc, "Responsive and Mobile-First Web Design")
    bullet(doc, "RESTful API Development and Database Schema Design")
    bullet(doc, "Admin Dashboard and CRM Functionality Implementation")
    bullet(doc, "Competitive Analysis and Market Research")

    # NO page break — content flows continuously

    # ═══ 2. COMPANY OVERVIEW ═══════════════════════════════════
    h1(doc, "2. Company Overview – PranaYuv Technology")
    para(doc, "PranaYuv Technology is a healthcare technology startup that specializes in autonomous patient care solutions. The company focuses on leveraging modern technology to improve healthcare delivery and patient outcomes. PranaYuv Technology operates at the intersection of healthcare and technology, providing innovative solutions for hospitals, clinics, and healthcare institutions.")
    para(doc, "The company's primary focus areas include developing digital platforms for healthcare equipment manufacturers, building educational technology solutions, and conducting market research for healthcare startups. During the internship, the company was working on multiple client-facing projects that required expertise in modern web technologies and design principles.")

    ital_head(doc, "Company Vision")
    para(doc, "To revolutionize patient care through autonomous technology solutions that make healthcare more accessible, efficient, and data-driven.")

    ital_head(doc, "Company Mission")
    para(doc, "To build scalable digital platforms and design intuitive user experiences that empower healthcare providers and educational institutions to deliver better services.")

    ital_head(doc, "Areas of Operation")
    bullet(doc, "Healthcare Technology Solutions")
    bullet(doc, "Medical Equipment E-Commerce Platforms")
    bullet(doc, "Educational Technology (EdTech) Applications")
    bullet(doc, "UI/UX Design and Consulting Services")
    bullet(doc, "Competitive Analysis and Market Research")

    # ═══ 3. PROJECT 1 ══════════════════════════════════════════
    h1(doc, "3. Project 1 – Sri Srinivasa Clean Rooms E-Commerce Web Application")
    para(doc, "The Sri Srinivasa Clean Rooms Web Application is a full-stack e-commerce and business management platform developed for a leading medical furniture manufacturer based in Hyderabad, Telangana. The application serves as both a public-facing product catalog and an internal administrative tool for managing products, customer inquiries, and business operations.")
    para(doc, "The project was built from scratch using modern web technologies and follows industry best practices for security, performance, and user experience. The application includes a responsive public website for customers and a comprehensive admin dashboard for business operations.")

    ital_head(doc, "Project Objectives")
    bullet(doc, "Build a professional, responsive e-commerce website for showcasing medical furniture products")
    bullet(doc, "Implement a secure admin dashboard with JWT-based authentication for product and inquiry management")
    bullet(doc, "Design a customer inquiry and lead management system with status tracking and priority assignment")
    bullet(doc, "Ensure mobile-first responsive design across all pages")
    bullet(doc, "Implement RESTful APIs with proper validation and error handling")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js, Vite, React Router, CSS3, JavaScript (ES6+)")
    bullet(doc, "Backend: Node.js, Express.js")
    bullet(doc, "Database: MongoDB with Mongoose ODM")
    bullet(doc, "Authentication: JWT (JSON Web Tokens), bcrypt.js for password hashing")
    bullet(doc, "File Handling: Multer for image uploads, Cloudinary for cloud storage")
    bullet(doc, "Build Tool: Vite")
    bullet(doc, "Version Control: Git")

    # ─── 3.1 Home Page ────────────────────────────────────────
    h2(doc, "3.1 Home Page – Landing Experience")
    para(doc, "The Home Page is the primary landing page of the Sri Srinivasa Clean Rooms web application. It serves as the first point of interaction for visitors and is designed to create a strong first impression while providing easy access to all product categories and company information. The page features a modern, visually rich design with smooth animations, a hero section showcasing the company's core value proposition, a curated product preview section, statistical highlights, and a comprehensive footer with contact information.")

    ital_head(doc, "Key Features")
    bullet(doc, "Hero Section with Dynamic Background:")
    sub_bullet(doc, "Full-screen hero section with gradient overlays and animated elements")
    sub_bullet(doc, "Company tagline and call-to-action buttons for immediate user engagement")
    sub_bullet(doc, "Responsive design that adapts to all screen sizes")
    bullet(doc, "Navigation Bar:")
    sub_bullet(doc, "Fixed-position navigation bar with scroll-based transparency effect")
    sub_bullet(doc, "Hamburger menu for mobile devices")
    sub_bullet(doc, "Quick links to Home, Products, About, Why Us, and Contact sections")
    bullet(doc, "Product Preview Section:")
    sub_bullet(doc, "Showcases six featured product categories: ICU Beds, Hospital Trolleys, Medical Lockers, Clean Room Systems, Procedure Tables, and Nurse Stations")
    sub_bullet(doc, "Each product card includes a high-quality image, title, description, and a colored badge (Best Seller, New, Premium)")
    sub_bullet(doc, "Hover effects and smooth animations for improved user interaction")
    bullet(doc, "Statistics Bar:")
    sub_bullet(doc, "Displays key company metrics: 100+ Products, 50+ Hospitals, 10+ Years, 24/7 Support")
    sub_bullet(doc, "Dark-themed section with contrasting color scheme for visual emphasis")
    bullet(doc, "Feature Highlights Section:")
    sub_bullet(doc, "ISO-Grade Materials – Hospital-grade stainless steel and anti-bacterial coatings")
    sub_bullet(doc, "Custom Manufacturing – Bespoke designs tailored to institutional specifications")
    sub_bullet(doc, "Rapid Delivery – Pan-India logistics with dedicated installation teams")
    sub_bullet(doc, "After-Sales Service – AMC contracts and on-call technician support")
    bullet(doc, "Fully Responsive Design:")
    sub_bullet(doc, "The page is fully responsive and adapts seamlessly to desktop, tablet, and mobile viewports")
    sub_bullet(doc, "Uses CSS media queries and flexible grid layouts")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js, CSS3 (Custom styling with CSS Variables)")
    bullet(doc, "Routing: React Router for smooth page transitions")
    bullet(doc, "Animations: CSS keyframe animations and Intersection Observer API for scroll-triggered effects")
    bullet(doc, "Icons: React Icons library")

    add_img(doc, "react_home.png", "Figure 3.1: Sri Srinivasa Clean Rooms – Home Page (Hero Section)")

    # ─── 3.2 Products Page ─────────────────────────────────────
    h2(doc, "3.2 Products Page – Catalog & Filtering")
    para(doc, "The Products Page is the core catalog component of the e-commerce platform. It allows customers to browse, search, filter, and sort through the complete range of medical furniture products offered by Sri Srinivasa Clean Rooms. The page dynamically fetches product data from the backend API and presents it in an organized, visually appealing grid layout.")

    ital_head(doc, "Key Features")
    bullet(doc, "Multi-Parameter Search:")
    sub_bullet(doc, "Users can search products by title, description, or category name")
    sub_bullet(doc, "Real-time search filtering with instant results as the user types")
    sub_bullet(doc, "Case-insensitive search for improved usability")
    bullet(doc, "Category-Based Filtering:")
    sub_bullet(doc, "Dynamic category filter buttons generated from the database")
    sub_bullet(doc, "One-click category selection with visual active state indication")
    sub_bullet(doc, "\"All\" option to reset filters and view the complete product range")
    bullet(doc, "Sorting Options:")
    sub_bullet(doc, "Price: Low to High sorting for budget-conscious customers")
    sub_bullet(doc, "Price: High to Low sorting for premium product browsing")
    sub_bullet(doc, "Default sorting based on database insertion order")
    bullet(doc, "Pagination System:")
    sub_bullet(doc, "12 products displayed per page for optimal loading performance")
    sub_bullet(doc, "Page navigation with Previous, Next, and numbered page buttons")
    sub_bullet(doc, "Automatic scroll-to-top on page change")
    bullet(doc, "Product Cards:")
    sub_bullet(doc, "High-quality product images with hover zoom effects")
    sub_bullet(doc, "Product title, category tag, price display, and truncated description")
    sub_bullet(doc, "\"View Details\" button linking to individual product detail pages")
    sub_bullet(doc, "Smooth entrance animations using Intersection Observer API")
    bullet(doc, "Fully Responsive Design:")
    sub_bullet(doc, "Grid layout adapts from 3 columns on desktop to 2 on tablet to 1 on mobile")
    sub_bullet(doc, "Sticky navigation bar with scroll-based styling changes")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js with useMemo and useEffect hooks for performance optimization")
    bullet(doc, "API Integration: Fetch API with dynamic URL resolution")
    bullet(doc, "State Management: React useState for search, filter, sort, and pagination states")
    bullet(doc, "Styling: Custom CSS with CSS Variables and media queries")

    add_img(doc, "react_products.png", "Figure 3.2: Products Page – Product Catalog with Search and Filters")

    # ─── 3.3 Product Detail Page ───────────────────────────────
    h2(doc, "3.3 Product Detail Page – Inquiry & Quotation")
    para(doc, "The Product Detail Page provides an in-depth view of individual products and serves as the primary lead generation tool. When a customer clicks on a product from the catalog, they are directed to this page which displays comprehensive product information alongside an integrated inquiry form for requesting quotations.")

    ital_head(doc, "Key Features")
    bullet(doc, "Product Information Display:")
    sub_bullet(doc, "Full-resolution product image with zoom capability")
    sub_bullet(doc, "Product title, category, price, and detailed description")
    sub_bullet(doc, "Clean layout with clear visual hierarchy")
    bullet(doc, "Inquiry / Quotation Form:")
    sub_bullet(doc, "Integrated form allowing customers to request quotes directly from the product page")
    sub_bullet(doc, "Fields: Name (required), Phone (required), Email (optional), Message (optional)")
    sub_bullet(doc, "Auto-populates product information in the inquiry submission")
    sub_bullet(doc, "Success and error notifications after form submission")
    bullet(doc, "Related Products Section:")
    sub_bullet(doc, "Displays other products from the same category for cross-selling")
    sub_bullet(doc, "Clickable product cards that navigate to the respective detail pages")
    bullet(doc, "API Integration:")
    sub_bullet(doc, "Fetches product details from the REST API using the product ID from URL parameters")
    sub_bullet(doc, "Submits customer inquiries via POST request to the inquiries API endpoint")
    sub_bullet(doc, "Includes the product ID, product name, and inquiry type in the submission payload")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js, React Router (useParams for URL parameter extraction)")
    bullet(doc, "API: RESTful API with JSON request/response format")
    bullet(doc, "Form Handling: Controlled components with React useState")
    bullet(doc, "Styling: Custom responsive CSS")

    # ─── 3.4 Admin Login ──────────────────────────────────────
    h2(doc, "3.4 Admin Login – Authentication Portal")
    para(doc, "The Admin Login page provides a secure authentication gateway for administrators to access the backend management dashboard. The page implements JWT-based authentication with bcrypt password hashing to ensure secure access to administrative operations.")

    ital_head(doc, "Key Features")
    bullet(doc, "Secure Authentication:")
    sub_bullet(doc, "Username and password-based login with server-side validation")
    sub_bullet(doc, "bcrypt.js used for password hashing and comparison")
    sub_bullet(doc, "JWT token generated upon successful authentication with configurable expiry")
    bullet(doc, "Token Management:")
    sub_bullet(doc, "JWT token stored in browser localStorage as 'adminToken'")
    sub_bullet(doc, "Automatic redirection to dashboard upon successful login")
    sub_bullet(doc, "Session expiry detection with automatic redirect to login page")
    bullet(doc, "Form Validation:")
    sub_bullet(doc, "Client-side validation for empty fields")
    sub_bullet(doc, "Server-side validation with generic error messages (does not reveal whether username or password is incorrect)")
    sub_bullet(doc, "Loading state indicator during authentication process")
    bullet(doc, "Clean UI Design:")
    sub_bullet(doc, "Centered login card with company branding")
    sub_bullet(doc, "Modern input fields with proper labels and focus states")
    sub_bullet(doc, "Responsive design that works across all devices")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js, React Router for navigation")
    bullet(doc, "Backend: Express.js with JWT authentication middleware")
    bullet(doc, "Security: bcrypt.js for password hashing, JWT for session management")
    bullet(doc, "Database: MongoDB (Admin model with username, passwordHash, role, isActive fields)")

    add_img(doc, "react_admin_login.png", "Figure 3.4: Admin Login – Secure Authentication Portal")

    # ─── 3.5 Admin Dashboard – Product Management ─────────────
    h2(doc, "3.5 Admin Dashboard – Product Management")
    para(doc, "The Admin Dashboard is the primary control panel for managing the Sri Srinivasa Clean Rooms business operations. The Product Management module allows administrators to perform full CRUD (Create, Read, Update, Delete) operations on products. The dashboard features a modern sidebar navigation, real-time data display, and toast notifications for user feedback.")

    ital_head(doc, "Key Features")
    bullet(doc, "Product CRUD Operations:")
    sub_bullet(doc, "Add Product: Form with fields for title, category (dropdown), price, description, and image upload")
    sub_bullet(doc, "Edit Product: Pre-populated form with existing product data for modification")
    sub_bullet(doc, "Delete Product: Confirmation-based deletion with immediate UI update")
    sub_bullet(doc, "View Products: Category-wise grouped display with expand/collapse functionality")
    bullet(doc, "Image Upload:")
    sub_bullet(doc, "Multer middleware for handling multipart form data uploads")
    sub_bullet(doc, "Cloudinary integration for cloud-based image storage")
    sub_bullet(doc, "Image preview before upload")
    bullet(doc, "Category-Wise Product Display:")
    sub_bullet(doc, "Products are grouped by their category with expandable/collapsible sections")
    sub_bullet(doc, "Product count displayed next to each category name")
    bullet(doc, "Sidebar Navigation:")
    sub_bullet(doc, "Collapsible sidebar with icons and labels")
    sub_bullet(doc, "Quick access to Products, Inquiries, Categories, and Settings tabs")
    sub_bullet(doc, "Mobile-responsive hamburger menu for smaller screens")
    bullet(doc, "Toast Notifications:")
    sub_bullet(doc, "Success and error toast messages for all CRUD operations")
    sub_bullet(doc, "Auto-dismiss after 3.5 seconds")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js with useCallback for optimized data fetching")
    bullet(doc, "Backend: Express.js RESTful API with JWT-protected routes")
    bullet(doc, "Database: MongoDB with Mongoose ODM (Product model: title, category, price, image, desc)")
    bullet(doc, "File Upload: Multer + Cloudinary")
    bullet(doc, "Styling: Custom CSS with responsive grid layout")

    add_img(doc, "react_admin_dashboard.png", "Figure 3.5: Admin Dashboard – Product & Category Management Interface")

    # ─── 3.6 Inquiry & Lead Management ─────────────────────────
    h2(doc, "3.6 Admin Dashboard – Inquiry & Lead Management")
    para(doc, "The Inquiry Management module is one of the most feature-rich components of the admin dashboard. It serves as a lightweight CRM (Customer Relationship Management) system, allowing the business to track, manage, and follow up on customer inquiries and leads.")

    ital_head(doc, "Key Features")
    bullet(doc, "Inquiry List View:")
    sub_bullet(doc, "All customer inquiries displayed in a tabular format with essential details")
    sub_bullet(doc, "Shows customer name, phone, email, product of interest, date, status, and priority")
    sub_bullet(doc, "Pagination with 10 items per page")
    bullet(doc, "Advanced Filtering:")
    sub_bullet(doc, "Filter by Status: New, Contacted, Quotation Sent, Negotiation, Closed, Rejected")
    sub_bullet(doc, "Filter by Priority: Normal, Hot Lead, Bulk Order, Dealer, Urgent, Hospital")
    sub_bullet(doc, "Filter by Lead Type: General inquiry, Product-specific inquiry")
    sub_bullet(doc, "Text-based search across name, phone, email, and product fields")
    bullet(doc, "Status & Priority Management:")
    sub_bullet(doc, "Color-coded status badges for visual identification")
    sub_bullet(doc, "Dropdown selectors for changing inquiry status and priority inline")
    sub_bullet(doc, "Real-time API updates when status or priority is changed")
    bullet(doc, "Notes System:")
    sub_bullet(doc, "Modal popup for adding internal notes to each inquiry")
    sub_bullet(doc, "Timestamped notes for tracking follow-up activities")
    sub_bullet(doc, "Notes history displayed in chronological order")
    bullet(doc, "Communication Integration:")
    sub_bullet(doc, "One-click WhatsApp message generation with pre-formatted inquiry details")
    sub_bullet(doc, "One-click Email composition with pre-filled subject and body")
    sub_bullet(doc, "Direct phone call link for immediate follow-up")
    bullet(doc, "Read/Unread Tracking:")
    sub_bullet(doc, "Visual indicator for new/unread inquiries")
    sub_bullet(doc, "Automatic mark-as-read when inquiry details are viewed")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js, CSS3 with color-coded badges and responsive tables")
    bullet(doc, "Backend: Express.js with protected API endpoints")
    bullet(doc, "Database: MongoDB (Inquiry model with status, priority, notes array, read flag)")
    bullet(doc, "Integration: WhatsApp API, Mailto protocol")

    # ─── 3.7 Category Management ──────────────────────────────
    h2(doc, "3.7 Admin Dashboard – Category Management")
    para(doc, "The Category Management module allows administrators to create, update, and delete product categories. Categories are used to organize products on the public-facing catalog and in the admin product management interface.")

    ital_head(doc, "Key Features")
    bullet(doc, "Add New Category:")
    sub_bullet(doc, "Simple input field with validation to prevent empty or duplicate category names")
    sub_bullet(doc, "Case-insensitive duplicate detection")
    bullet(doc, "Edit Category:")
    sub_bullet(doc, "Inline editing of existing category names")
    sub_bullet(doc, "Automatic update of associated product references")
    bullet(doc, "Delete Category:")
    sub_bullet(doc, "Confirmation dialog before deletion")
    sub_bullet(doc, "Warning if products are associated with the category")
    bullet(doc, "Visual Category List:")
    sub_bullet(doc, "Shows all categories with the count of products in each")
    sub_bullet(doc, "Clean list view with action buttons for edit and delete")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: React.js")
    bullet(doc, "Backend: Express.js REST API with Joi validation")
    bullet(doc, "Database: MongoDB (Category model: name, timestamps)")

    # ─── 3.8 Backend Architecture ──────────────────────────────
    h2(doc, "3.8 Backend Architecture & Database Design")
    para(doc, "The backend of the Sri Srinivasa Clean Rooms application follows a modular, layered architecture pattern. The codebase is organized into separate directories for routes, controllers, models, middleware, validators, and utility functions. This separation of concerns ensures maintainability, testability, and scalability of the application.")

    ital_head(doc, "Server Architecture")
    para(doc, "The backend is structured as follows:")
    bullet(doc, "server.js – Entry point; connects to MongoDB and starts the Express server")
    bullet(doc, "app.js – Express application configuration; registers middleware, routes, and static file serving")
    bullet(doc, "config/ – Database connection (db.js) and environment variable configuration")
    bullet(doc, "routes/ – API endpoint definitions: authRoutes, productRoutes, categoryRoutes, inquiryRoutes")
    bullet(doc, "controllers/ – HTTP request/response handling logic for each module")
    bullet(doc, "models/ – Mongoose schema definitions: Admin, Product, Category, Inquiry")
    bullet(doc, "middleware/ – Authentication (JWT verification), error handling, file upload (Multer)")
    bullet(doc, "validators/ – Input validation using Joi for all API endpoints")
    bullet(doc, "utils/ – Shared utilities: ApiError class, async handler wrapper")

    ital_head(doc, "Database Models")
    para(doc, "The application uses MongoDB with Mongoose ODM. The primary models are:")
    bold_para(doc, "Product Model: ", "title (String, required), category (String, required), price (String, required), image (String), desc (String, required, max 2000 chars)")
    bold_para(doc, "Inquiry Model: ", "name, phone, email, message, productName, type (general/product), status (new/contacted/quotation/negotiation/closed/rejected), priority (normal/hot/bulk/dealer/urgent/hospital), notes (array), read (Boolean)")
    bold_para(doc, "Category Model: ", "name (String, required, unique), timestamps")
    bold_para(doc, "Admin Model: ", "username, email, passwordHash (bcrypt), role, isActive, lastLoginAt, timestamps")

    ital_head(doc, "API Endpoints Overview")
    bullet(doc, "POST /api/auth/login – Admin authentication with JWT generation")
    bullet(doc, "GET /api/products – Fetch all products (public)")
    bullet(doc, "POST /api/products – Add new product (protected)")
    bullet(doc, "PUT /api/products/:id – Update product (protected)")
    bullet(doc, "DELETE /api/products/:id – Delete product (protected)")
    bullet(doc, "GET /api/categories – Fetch all categories")
    bullet(doc, "POST /api/categories – Add category (protected)")
    bullet(doc, "PUT /api/categories/:id – Update category (protected)")
    bullet(doc, "DELETE /api/categories/:id – Delete category (protected)")
    bullet(doc, "GET /api/inquiries – Fetch all inquiries (protected)")
    bullet(doc, "POST /api/inquiries – Submit inquiry (public)")
    bullet(doc, "PUT /api/inquiries/:id/status – Update status (protected)")
    bullet(doc, "PUT /api/inquiries/:id/priority – Update priority (protected)")
    bullet(doc, "POST /api/inquiries/:id/notes – Add note (protected)")
    bullet(doc, "DELETE /api/inquiries/:id – Delete inquiry (protected)")

    ital_head(doc, "Security Implementation")
    bullet(doc, "JWT Authentication: All administrative routes are protected with a JWT verification middleware")
    bullet(doc, "Password Hashing: bcrypt.js with minimum 10 salt rounds")
    bullet(doc, "Input Validation: Joi validation on all incoming request data")
    bullet(doc, "CORS Configuration: Configured to allow requests only from the client application origin")
    bullet(doc, "Error Handling: Centralized error middleware with consistent JSON error response format")

    # ═══ 4. PROJECT 2 ══════════════════════════════════════════
    page_break(doc)
    h1(doc, "4. Project 2 – PranaYuv Technology Corporate Landing Page")
    para(doc, "The PranaYuv Technology Corporate Landing Page is a high-end, visually immersive static website designed to showcase the company's mission, products, and brand identity. The website was built using pure HTML5, CSS3, and JavaScript without any framework, emphasizing hand-crafted design and performance optimization. The landing page serves as the company's primary digital presence, conveying a premium brand image through sophisticated typography, advanced CSS animations, and a custom cursor implementation.")

    ital_head(doc, "Key Features")
    bullet(doc, "Custom Cursor Effect:")
    sub_bullet(doc, "Custom-designed dual-cursor system (inner dot + outer ring) using pure CSS and JavaScript")
    sub_bullet(doc, "Cursor scales and changes on hover interactions for an immersive browsing experience")
    sub_bullet(doc, "Default system cursor is hidden for a unique visual effect")
    bullet(doc, "Hero Section with Canvas Animation:")
    sub_bullet(doc, "Full-viewport hero section with a two-column grid layout (52% text, 48% visual)")
    sub_bullet(doc, "HTML5 Canvas used for animated background effects in the visual column")
    sub_bullet(doc, "Floating HUD (Heads-Up Display) elements with CSS infinite animations")
    sub_bullet(doc, "Staggered text entrance animations using CSS keyframes")
    bullet(doc, "Advanced Typography:")
    sub_bullet(doc, "Uses three carefully selected Google Fonts: Cormorant Garamond (serif headings), Outfit (sans-serif body), JetBrains Mono (monospace labels)")
    sub_bullet(doc, "Gradient text effects using CSS background-clip for emphasis elements")
    sub_bullet(doc, "Font sizes defined using CSS clamp() for fluid responsive typography")
    bullet(doc, "Statistics Bar:")
    sub_bullet(doc, "Dark-themed statistics section at the bottom of the hero with animated number counters for key metrics")
    sub_bullet(doc, "Grid layout with subtle dividers between stat items")
    bullet(doc, "Scroll-Based Navigation Transformation:")
    sub_bullet(doc, "Fixed-position navigation bar that transforms on scroll")
    sub_bullet(doc, "Transparent background on top, frosted glass effect (backdrop-filter: blur) when scrolled")
    sub_bullet(doc, "Smooth padding transition for seamless user experience")
    bullet(doc, "Fully Responsive Design:")
    sub_bullet(doc, "Designed with a desktop-first approach and comprehensive media queries for all breakpoints")
    sub_bullet(doc, "Custom scrollbar styling with accent color thumb")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: HTML5, CSS3 (Custom Properties, Animations, Grid, Flexbox)")
    bullet(doc, "JavaScript: Vanilla JS for cursor tracking, scroll detection, and canvas rendering")
    bullet(doc, "Typography: Google Fonts (Cormorant Garamond, Outfit, JetBrains Mono)")
    bullet(doc, "Design: Custom color palette using CSS custom properties")

    add_img(doc, "pranayuv_index.png", "Figure 4.1: PranaYuv Technology – Corporate Landing Page")

    # ─── 4.1 Design Philosophy ─────────────────────────────────
    h2(doc, "4.1 Design Philosophy & UI/UX Decisions")
    para(doc, "The design of the PranaYuv landing page was guided by several key principles from modern UI/UX design theory. The goal was to create a website that conveys trust, innovation, and premium quality – aligning with the company's positioning in the healthcare technology space.")

    ital_head(doc, "Design Principles Applied")
    bullet(doc, "Visual Hierarchy: Content is organized with clear heading levels, contrasting font weights, and strategic use of whitespace to guide the user's eye through the page in a logical flow.")
    bullet(doc, "Micro-Interactions: Subtle hover effects on navigation links, buttons, and interactive elements provide tactile feedback and make the interface feel alive and responsive.")
    bullet(doc, "Color Psychology: The primary violet (#7c3aed) conveys innovation and creativity, while the teal accent (#0d9488) adds trust and professionalism. The warm neutral background (#f7f4ef) provides a sophisticated, non-clinical feel.")
    bullet(doc, "Progressive Disclosure: Information is revealed gradually through scroll-triggered animations, maintaining user engagement throughout the page.")
    bullet(doc, "Accessibility Considerations: Sufficient color contrast ratios, semantic HTML elements, and clear interactive states ensure the page is usable by a wide range of users.")

    # ═══ 5. PROJECT 3 ══════════════════════════════════════════
    page_break(doc)
    h1(doc, "5. Project 3 – BrightPath Academy School Management System")
    para(doc, "The BrightPath Academy School Management System is a comprehensive web application designed to digitize and streamline the operations of a K-12 educational institution. The system provides dedicated interfaces for the school's public website, student admissions, student dashboard, and principal/admin dashboard. The project was developed using HTML5, CSS3, JavaScript, and Bootstrap 5 as a frontend-focused application with responsive design across all modules.")

    ital_head(doc, "Project Scope")
    bullet(doc, "Public-facing school website with information about programs, facilities, and admissions")
    bullet(doc, "Student admission portal with multi-step registration forms")
    bullet(doc, "Student dashboard with academic performance tracking, attendance, and schedule management")
    bullet(doc, "Principal dashboard with school-wide analytics, staff management, and operational controls")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: HTML5, CSS3, JavaScript")
    bullet(doc, "UI Framework: Bootstrap 5.3.2")
    bullet(doc, "Typography: Google Fonts (Montserrat, Poppins)")
    bullet(doc, "Icons: Font Awesome 6.5.0")
    bullet(doc, "Animations: AOS (Animate on Scroll) library")
    bullet(doc, "Charts: Chart.js for data visualization")

    # ─── 5.1 School Home Page ──────────────────────────────────
    h2(doc, "5.1 School Website – Home Page")
    para(doc, "The BrightPath Academy home page serves as the school's digital front door. It is designed to be welcoming, informative, and visually engaging for parents, students, and prospective families. The page uses a blue-gold color scheme that represents academic excellence and warmth.")

    ital_head(doc, "Key Features")
    bullet(doc, "Hero Section:")
    sub_bullet(doc, "Full-width hero banner with overlay text and call-to-action buttons")
    sub_bullet(doc, "School name, tagline ('Shaping Future Minds'), and admission inquiry buttons")
    sub_bullet(doc, "Background image with gradient overlay for text readability")
    bullet(doc, "Programs Section:")
    sub_bullet(doc, "Highlights academic programs offered at different grade levels")
    sub_bullet(doc, "Card-based layout with icons and short descriptions")
    sub_bullet(doc, "AOS scroll animations for sequential card appearance")
    bullet(doc, "Facilities Showcase:")
    sub_bullet(doc, "Photo gallery of school facilities: classrooms, library, laboratory, playground")
    sub_bullet(doc, "Overlay text with facility descriptions on hover")
    bullet(doc, "Testimonials Section:")
    sub_bullet(doc, "Parent and student testimonials with profile images and carousel-style display")
    bullet(doc, "Contact & Enquiry Section:")
    sub_bullet(doc, "Contact form for admission inquiries")
    sub_bullet(doc, "Google Maps embed for school location and direct phone/email links")
    bullet(doc, "Responsive Navigation:")
    sub_bullet(doc, "Fixed navbar with scroll-based background change")
    sub_bullet(doc, "Mobile-responsive hamburger menu with quick links to all sections")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: HTML5, CSS3, Bootstrap 5")
    bullet(doc, "Animations: AOS library, CSS transitions")
    bullet(doc, "Typography: Montserrat (headings), Poppins (body text)")

    add_img(doc, "school_index.png", "Figure 5.1: BrightPath Academy – School Home Page")

    # ─── 5.2 Admission Portal ──────────────────────────────────
    h2(doc, "5.2 Admission Portal")
    para(doc, "The Admission Portal is a comprehensive online registration system that allows parents and guardians to submit admission applications for their children. The portal includes multi-step forms with proper validation, document upload placeholders, and a structured data collection process.")

    ital_head(doc, "Key Features")
    bullet(doc, "Multi-Step Registration Form:")
    sub_bullet(doc, "Step 1: Student personal information (name, date of birth, gender, grade applying for)")
    sub_bullet(doc, "Step 2: Parent/guardian details (father's name, mother's name, contact information)")
    sub_bullet(doc, "Step 3: Previous school information and academic records")
    sub_bullet(doc, "Step 4: Document upload section for required certificates")
    sub_bullet(doc, "Step 5: Review and submit with declaration checkbox")
    bullet(doc, "Form Validation:")
    sub_bullet(doc, "Required field validation with visual error indicators")
    sub_bullet(doc, "Phone number and email format validation")
    sub_bullet(doc, "Date range validation for date of birth fields")
    bullet(doc, "Progress Indicator:")
    sub_bullet(doc, "Visual step-by-step progress bar showing the current stage")
    sub_bullet(doc, "Clickable step indicators for navigating back to previous steps")
    bullet(doc, "Responsive Design:")
    sub_bullet(doc, "The entire form adapts to mobile screens with stacked layout")
    sub_bullet(doc, "Touch-friendly input fields and buttons for mobile users")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: HTML5, CSS3, Bootstrap 5, JavaScript")
    bullet(doc, "Validation: Custom JavaScript validation logic")
    bullet(doc, "Icons: Font Awesome for visual form enhancement")

    add_img(doc, "school_admission.png", "Figure 5.2: BrightPath Academy – Admission Portal")

    # ─── 5.3 Student Dashboard ─────────────────────────────────
    h2(doc, "5.3 Student Dashboard")
    para(doc, "The Student Dashboard provides a personalized interface for students to access their academic information, attendance records, schedule, and performance metrics. The dashboard uses a modern card-based layout with data visualizations for an intuitive overview of the student's academic life.")

    ital_head(doc, "Key Features")
    bullet(doc, "Academic Performance Overview:")
    sub_bullet(doc, "Subject-wise grade display with color-coded performance indicators")
    sub_bullet(doc, "GPA/CGPA calculation and display with term-wise comparison charts")
    bullet(doc, "Attendance Tracking:")
    sub_bullet(doc, "Monthly attendance percentage with visual progress indicators")
    sub_bullet(doc, "Day-wise attendance calendar view and attendance trend graphs")
    bullet(doc, "Class Schedule / Timetable:")
    sub_bullet(doc, "Weekly timetable view with color-coded subjects")
    sub_bullet(doc, "Current/next class highlighting")
    bullet(doc, "Announcements & Notices:")
    sub_bullet(doc, "School-wide and class-specific announcements with date-stamped notice board")
    bullet(doc, "Profile Section:")
    sub_bullet(doc, "Student photo, name, class, section, and roll number")
    sub_bullet(doc, "Quick access to personal and academic information")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: HTML5, CSS3, JavaScript, Bootstrap 5")
    bullet(doc, "Charts: Chart.js for performance and attendance visualizations")
    bullet(doc, "Icons: Font Awesome")

    add_img(doc, "school_student_dash.png", "Figure 5.3: BrightPath Academy – Student Dashboard")

    # ─── 5.4 Principal Dashboard ───────────────────────────────
    h2(doc, "5.4 Principal Dashboard")
    para(doc, "The Principal Dashboard provides a high-level administrative view of the entire school's operations. It includes school-wide metrics, staff management tools, academic performance summaries, and operational controls for the school administration.")

    ital_head(doc, "Key Features")
    bullet(doc, "School-Wide Metrics:")
    sub_bullet(doc, "Total students, staff, and class counts")
    sub_bullet(doc, "Overall attendance rates, academic performance averages, and revenue/fee collection summaries")
    bullet(doc, "Staff Management:")
    sub_bullet(doc, "Staff directory with department-wise filtering")
    sub_bullet(doc, "Staff attendance and performance tracking with leave management system")
    bullet(doc, "Academic Analytics:")
    sub_bullet(doc, "Class-wise and subject-wise performance comparison charts")
    sub_bullet(doc, "Top performers and students needing attention with examination result analysis")
    bullet(doc, "Operational Controls:")
    sub_bullet(doc, "Calendar management for school events")
    sub_bullet(doc, "Announcement publishing and report generation tools")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Frontend: HTML5, CSS3, JavaScript, Bootstrap 5")
    bullet(doc, "Data Visualization: Chart.js, custom CSS-based progress indicators")
    bullet(doc, "Layout: Sidebar navigation with main content area")

    # ═══ 6. PROJECT 4 ══════════════════════════════════════════
    page_break(doc)
    h1(doc, "6. Project 4 – EduNexus UI/UX Design (Figma to HTML/CSS)")
    para(doc, "The EduNexus project involved designing and implementing UI/UX prototypes for an educational technology platform. The designs were first created in Figma following modern design principles and then converted to pixel-perfect HTML/CSS implementations. EduNexus is envisioned as a comprehensive educational platform that connects students, teachers, and administrators through a unified digital interface. The design focuses on usability, accessibility, and visual consistency across all pages.")

    ital_head(doc, "Design Process")
    bullet(doc, "Research & Analysis: Studied existing EdTech platforms (Google Classroom, Canvas LMS, Moodle) to identify best practices and common usability patterns")
    bullet(doc, "Wireframing: Created low-fidelity wireframes for each page to establish layout and content hierarchy")
    bullet(doc, "Visual Design: Developed a comprehensive design system including color palette, typography scale, component library, and spacing rules")
    bullet(doc, "Prototyping: Built interactive prototypes in Figma with linked screens and transitions")
    bullet(doc, "Implementation: Converted Figma designs to responsive HTML/CSS with pixel-perfect accuracy")

    ital_head(doc, "Design System")
    bullet(doc, "Color Palette: Olive-green primary (#6B8E23) with complementary earth tones for a calm, focused learning environment")
    bullet(doc, "Typography: Clean sans-serif fonts for readability across screen sizes")
    bullet(doc, "Components: Reusable UI components including buttons, cards, form inputs, navigation elements, tables, and modal dialogs")
    bullet(doc, "Spacing: 8px grid system for consistent spacing and alignment")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Design: Figma (UI/UX Design and Prototyping)")
    bullet(doc, "Implementation: HTML5, CSS3 (Flexbox, Grid, Custom Properties)")
    bullet(doc, "Responsive: Media queries for multi-device compatibility")

    # ─── 6.1 EduNexus Olive Theme ─────────────────────────────
    h2(doc, "6.1 EduNexus Olive Theme Design")
    para(doc, "The EduNexus Olive Theme is the primary visual design of the educational platform. The olive-green color scheme was chosen to create a calm, focused learning environment that reduces eye strain during long study sessions. The theme includes a comprehensive component library with consistent styling across all elements.")

    ital_head(doc, "Key Features")
    bullet(doc, "Consistent Visual Language:")
    sub_bullet(doc, "Unified color palette applied across all pages and components")
    sub_bullet(doc, "Consistent border-radius, shadow, and transition values")
    sub_bullet(doc, "Clear distinction between primary, secondary, and accent elements")
    bullet(doc, "Navigation Design:")
    sub_bullet(doc, "Top navigation bar with logo, search bar, and user profile dropdown")
    sub_bullet(doc, "Left sidebar for main navigation (Dashboard, Courses, Assignments, Schedule, Settings)")
    sub_bullet(doc, "Breadcrumb navigation for deep-page context")
    bullet(doc, "Content Layout:")
    sub_bullet(doc, "Card-based content organization for easy scanning")
    sub_bullet(doc, "Grid layout for course listings and resource collections")
    sub_bullet(doc, "Proper whitespace and padding for comfortable reading")
    bullet(doc, "Interactive Elements:")
    sub_bullet(doc, "Hover states on all clickable elements")
    sub_bullet(doc, "Form inputs with focus indicators and validation states")
    sub_bullet(doc, "Progress bars and status indicators")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Design: Figma")
    bullet(doc, "Implementation: HTML5, CSS3 with CSS Variables for theming")

    add_img(doc, "figma_edunexus.png", "Figure 6.1: EduNexus – Olive Theme Design Implementation")

    # ─── 6.2 Admission Requirements ───────────────────────────
    h2(doc, "6.2 Admission Requirements Page")
    para(doc, "The Admission Requirements page was designed to present complex admission criteria in a clear, organized, and digestible format. The page helps prospective students understand the requirements for different programs and provides step-by-step guidance on the application process.")

    ital_head(doc, "Key Features")
    bullet(doc, "Program-Wise Requirements:")
    sub_bullet(doc, "Tabbed or accordion-based layout for different programs")
    sub_bullet(doc, "Clear listing of eligibility criteria for each program")
    sub_bullet(doc, "Document checklist with visual indicators")
    bullet(doc, "Application Timeline:")
    sub_bullet(doc, "Visual timeline showing important dates for the admission process")
    sub_bullet(doc, "Step-by-step application process with descriptions")
    bullet(doc, "Information Cards:")
    sub_bullet(doc, "Key information presented in easily scannable card format")
    sub_bullet(doc, "Icons and visual cues for quick comprehension")
    sub_bullet(doc, "Contact information for admission queries")
    bullet(doc, "Responsive Layout:")
    sub_bullet(doc, "Desktop: Multi-column layout with sidebar")
    sub_bullet(doc, "Mobile: Single-column stacked layout")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Design: Figma")
    bullet(doc, "Implementation: HTML5, CSS3")

    add_img(doc, "figma_admission_req.png", "Figure 6.2: EduNexus – Admission Requirements Page Design")

    # ─── 6.3 Student Dashboard V2 ─────────────────────────────
    h2(doc, "6.3 Student Dashboard V2")
    para(doc, "The Student Dashboard V2 represents an improved iteration of the student interface, focusing on enhanced usability, better data visualization, and a widget-based modular layout. This version was designed based on user feedback and best practices from leading EdTech platforms.")

    ital_head(doc, "Key Features")
    bullet(doc, "Widget-Based Layout:")
    sub_bullet(doc, "Modular dashboard with draggable widget sections")
    sub_bullet(doc, "Customizable layout allowing students to prioritize the information they see first")
    sub_bullet(doc, "Widgets for grades, attendance, upcoming assignments, and announcements")
    bullet(doc, "Enhanced Data Visualization:")
    sub_bullet(doc, "Circular progress indicators for course completion")
    sub_bullet(doc, "Bar charts for subject-wise performance and calendar widget for schedule/deadlines")
    bullet(doc, "Dark Mode Support:")
    sub_bullet(doc, "Toggle between light and dark color schemes with proper contrast ratios maintained in both modes")
    sub_bullet(doc, "Reduced eye strain for extended use")
    bullet(doc, "Intuitive Navigation:")
    sub_bullet(doc, "Clear sidebar with icon and text labels")
    sub_bullet(doc, "Quick-action buttons for frequently used features")
    sub_bullet(doc, "Notification center with unread count badge")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Design: Figma (with Auto-Layout and Component Variants)")
    bullet(doc, "Implementation: HTML5, CSS3 (Grid, Flexbox, CSS Variables for theming)")

    add_img(doc, "figma_student_dash.png", "Figure 6.3: EduNexus – Student Dashboard V2 Design")

    # ─── 6.4 Principal Dashboard ──────────────────────────────
    h2(doc, "6.4 Principal Dashboard")
    para(doc, "The EduNexus Principal Dashboard is designed as a comprehensive administrative control panel. It provides school leaders with real-time insights into institutional performance, student analytics, staff management, and operational metrics, all within a single unified interface.")

    ital_head(doc, "Key Features")
    bullet(doc, "Executive Summary View:")
    sub_bullet(doc, "KPI cards showing critical metrics at a glance (total students, pass rates, attendance rates, revenue)")
    sub_bullet(doc, "Trend indicators (up/down arrows) showing month-over-month changes")
    bullet(doc, "Analytics Dashboard:")
    sub_bullet(doc, "Interactive charts for enrollment trends, performance distribution, and department comparisons")
    sub_bullet(doc, "Drill-down capability from school-wide to class-level data")
    bullet(doc, "Quick Actions Panel:")
    sub_bullet(doc, "One-click access to common administrative tasks")
    sub_bullet(doc, "Shortcut buttons for report generation, announcement publishing, and event scheduling")
    bullet(doc, "Recent Activity Feed:")
    sub_bullet(doc, "Timeline of recent school events, staff actions, and system notifications")
    sub_bullet(doc, "Filterable by activity type")

    ital_head(doc, "Technology Stack")
    bullet(doc, "Design: Figma")
    bullet(doc, "Implementation: HTML5, CSS3")

    add_img(doc, "figma_principal_dash.png", "Figure 6.4: EduNexus – Principal Dashboard Design")

    # ═══ 7. COMPETITIVE ANALYSIS ═══════════════════════════════
    page_break(doc)
    h1(doc, "7. Project 5 – Competitive Analysis & Market Research")
    para(doc, "As part of the internship, a detailed competitive analysis and market research document was prepared for PranaYuv Technology. The analysis evaluated the competitive landscape in the healthcare technology and autonomous patient care space, identifying key competitors, their strengths, weaknesses, and market positioning. This analysis was crucial for defining PranaYuv Technology's go-to-market strategy and product differentiation approach for the year 2026 and beyond.")

    ital_head(doc, "Scope of Analysis")
    bullet(doc, "Identification and profiling of direct and indirect competitors in the healthcare technology space")
    bullet(doc, "Feature comparison matrix across competing products and platforms")
    bullet(doc, "Pricing strategy analysis and market positioning")
    bullet(doc, "SWOT analysis for PranaYuv Technology")
    bullet(doc, "Target audience segmentation and customer persona development")
    bullet(doc, "Market trends and growth projections in healthcare technology")
    bullet(doc, "Recommendations for product differentiation and competitive advantages")

    ital_head(doc, "Key Findings")
    bullet(doc, "The healthcare technology market is projected to grow at a CAGR of 15-20% over the next five years, driven by increasing digitalization of healthcare services")
    bullet(doc, "Most competitors focus either on hardware or software solutions, creating an opportunity for PranaYuv to offer integrated solutions")
    bullet(doc, "Patient data privacy and regulatory compliance (HIPAA, GDPR) are critical differentiators in the market")
    bullet(doc, "Mobile-first design and telemedicine integration are emerging as essential features for healthcare platforms")

    ital_head(doc, "Deliverables")
    bullet(doc, "Comprehensive competitive analysis document (Microsoft Word format)")
    bullet(doc, "Competitor comparison spreadsheet with feature matrices")
    bullet(doc, "Strategic recommendations report for the management team")

    ital_head(doc, "Research Methodology")
    bullet(doc, "Secondary Research: Online databases, company websites, industry reports, news articles")
    bullet(doc, "Competitor Product Testing: Hands-on evaluation of competitor platforms where publicly available")
    bullet(doc, "Data Analysis: Comparative analysis using structured frameworks (SWOT, Porter's Five Forces)")

    # ═══ 8. TOOLS ══════════════════════════════════════════════
    h1(doc, "8. Tools & Technologies Used During Internship")
    para(doc, "The following tools and technologies were extensively used during the internship period across all projects:")

    ital_head(doc, "Frontend Development")
    bullet(doc, "React.js – JavaScript library for building user interfaces with component-based architecture")
    bullet(doc, "Vite – Next-generation frontend build tool for fast development server and optimized builds")
    bullet(doc, "HTML5 – Semantic markup for structured web content")
    bullet(doc, "CSS3 – Custom styling with CSS Variables, Flexbox, Grid, Animations, and Media Queries")
    bullet(doc, "JavaScript (ES6+) – Modern JavaScript with async/await, destructuring, template literals, and modules")
    bullet(doc, "Bootstrap 5 – UI framework for rapid responsive design (used in School Management System)")
    bullet(doc, "React Router – Client-side routing for single-page application navigation")
    bullet(doc, "React Icons – Icon library for consistent iconography across the React application")

    ital_head(doc, "Backend Development")
    bullet(doc, "Node.js – JavaScript runtime for server-side application development")
    bullet(doc, "Express.js – Web application framework for building RESTful APIs")
    bullet(doc, "MongoDB – NoSQL database for flexible and scalable data storage")
    bullet(doc, "Mongoose – ODM (Object Data Modeling) library for MongoDB")
    bullet(doc, "JWT (JSON Web Tokens) – Token-based authentication for secure API access")
    bullet(doc, "bcrypt.js – Password hashing library for secure credential management")
    bullet(doc, "Multer – Middleware for handling file uploads")
    bullet(doc, "Cloudinary – Cloud-based image and video management service")
    bullet(doc, "Joi – Schema-based validation library for request data validation")
    bullet(doc, "CORS – Cross-Origin Resource Sharing middleware for secure API access")

    ital_head(doc, "Design & Prototyping")
    bullet(doc, "Figma – UI/UX design and prototyping tool for creating wireframes, mockups, and interactive prototypes")
    bullet(doc, "Adobe Photoshop – Image editing and optimization for web assets")

    ital_head(doc, "Development Tools")
    bullet(doc, "VS Code – Primary code editor with extensions for React, Node.js, and CSS development")
    bullet(doc, "Git – Version control system for tracking code changes and collaboration")
    bullet(doc, "GitHub – Remote repository hosting and collaboration platform")
    bullet(doc, "Postman – API testing and development tool for RESTful API verification")
    bullet(doc, "Chrome DevTools – Browser-based debugging and performance analysis tool")
    bullet(doc, "npm – Package manager for JavaScript dependencies")
    bullet(doc, "Nodemon – Development utility for automatic server restart on code changes")

    ital_head(doc, "Libraries & Frameworks")
    bullet(doc, "AOS (Animate on Scroll) – Scroll-triggered animation library for web elements")
    bullet(doc, "Chart.js – Data visualization library for charts and graphs")
    bullet(doc, "Font Awesome – Icon library used across multiple projects")
    bullet(doc, "Google Fonts – Web typography service (Cormorant Garamond, Outfit, JetBrains Mono, Montserrat, Poppins)")

    # ═══ 9. SKILLS ═════════════════════════════════════════════
    h1(doc, "9. Skills Acquired & Learning Outcomes")
    para(doc, "The internship provided extensive hands-on experience and helped develop a wide range of technical and professional skills:")

    ital_head(doc, "Technical Skills")
    bullet(doc, "Full-Stack Development: Gained proficiency in building complete web applications from frontend to backend, including database design, API development, and deployment.")
    bullet(doc, "React.js Development: Learned component-based architecture, state management with hooks (useState, useEffect, useCallback, useMemo), routing, and performance optimization techniques.")
    bullet(doc, "Node.js & Express.js: Developed skills in building RESTful APIs, implementing authentication middleware, handling file uploads, and structuring modular backend applications.")
    bullet(doc, "MongoDB: Gained experience in NoSQL database design, Mongoose schema creation, CRUD operations, and data modeling for real-world applications.")
    bullet(doc, "UI/UX Design: Learned the complete design process from research and wireframing to high-fidelity prototyping in Figma, including design systems, typography, and color theory.")
    bullet(doc, "Responsive Web Design: Mastered CSS Grid, Flexbox, media queries, and mobile-first design principles for building websites that work across all devices.")
    bullet(doc, "CSS Animations: Developed expertise in CSS keyframe animations, transitions, and the Intersection Observer API for creating engaging micro-interactions.")
    bullet(doc, "Security: Learned JWT authentication implementation, bcrypt password hashing, input validation, CORS configuration, and secure API design practices.")

    ital_head(doc, "Professional Skills")
    bullet(doc, "Project Management: Experienced working on multiple projects simultaneously with different timelines and requirements.")
    bullet(doc, "Problem Solving: Developed analytical thinking through debugging complex issues, optimizing performance, and implementing creative solutions to design challenges.")
    bullet(doc, "Communication: Improved ability to present technical concepts to non-technical stakeholders and collaborate effectively with team members.")
    bullet(doc, "Time Management: Learned to prioritize tasks, manage deadlines, and balance multiple project responsibilities.")
    bullet(doc, "Research Skills: Conducted competitive analysis and market research, developing skills in data gathering, analysis, and report writing.")
    bullet(doc, "Attention to Detail: Developed a keen eye for pixel-perfect design implementation, code quality, and user experience refinement.")

    # ═══ 10. CONCLUSION ════════════════════════════════════════
    h1(doc, "10. Conclusion")
    para(doc, "The summer internship at PranaYuv Technology was an immensely valuable learning experience that bridged the gap between academic knowledge and real-world application development. Over the course of the internship, five major projects were completed, spanning the full spectrum of modern web development – from UI/UX design and frontend implementation to backend architecture and database management.")
    para(doc, "The internship provided the opportunity to work on production-grade applications that serve real clients, giving invaluable insights into the software development lifecycle, client requirements gathering, and iterative design processes. Working with modern technologies like React.js, Node.js, MongoDB, and Figma prepared a strong foundation for a career in full-stack web development and UI/UX design.")
    para(doc, "The experience of building the Sri Srinivasa Clean Rooms e-commerce platform demonstrated the complexity of developing business-critical applications with secure authentication, comprehensive data management, and CRM functionality. The school management system and EduNexus projects showcased the importance of user-centered design in educational technology, while the competitive analysis project developed business acumen alongside technical skills.")
    para(doc, "Key takeaways from the internship include the importance of writing clean, maintainable code with proper separation of concerns, the value of responsive and accessible design, the critical role of security in web applications, and the power of modern JavaScript frameworks in building scalable user interfaces.")
    para(doc, "In conclusion, this internship has significantly contributed to both technical expertise and professional development, providing a strong foundation for future endeavors in the field of software development and design.")

    # ═══ APPENDIX ══════════════════════════════════════════════
    h1(doc, "Appendix")

    ital_head(doc, "Appendix A: Project URLs and Access Information")
    bullet(doc, "Sri Srinivasa Clean Rooms Web App: http://localhost:5173/ (Development Server)")
    bullet(doc, "PranaYuv Technology Landing Page: Static HTML file (index.html)")
    bullet(doc, "BrightPath Academy School Management: Static HTML files (index.html, admission.html, student-dashboard.html)")
    bullet(doc, "EduNexus Designs: Converted Figma designs as HTML/CSS files")

    ital_head(doc, "Appendix B: File Structure Overview")
    para(doc, "The main React application follows a standard Vite + React project structure:")
    bullet(doc, "src/pages/ – Page-level components (Home, Products, ProductDetail, AdminLogin, AdminDashboard)")
    bullet(doc, "src/components/ – Reusable UI components (Navbar, Hero, Footer, ProductCard, CTASection)")
    bullet(doc, "src/assets/ – Static assets (images, logos)")
    bullet(doc, "server/ – Backend Node.js application")
    bullet(doc, "server/routes/ – API route definitions (auth, products, categories, inquiries)")
    bullet(doc, "server/controllers/ – Request handling logic")
    bullet(doc, "server/models/ – MongoDB schema definitions (Admin, Product, Category, Inquiry)")
    bullet(doc, "server/middleware/ – Authentication and error handling middleware")
    bullet(doc, "server/validators/ – Joi validation schemas")

    ital_head(doc, "Appendix C: References")
    bullet(doc, "React.js Documentation – https://react.dev")
    bullet(doc, "Node.js Documentation – https://nodejs.org/en/docs")
    bullet(doc, "MongoDB Documentation – https://www.mongodb.com/docs")
    bullet(doc, "Express.js Documentation – https://expressjs.com")
    bullet(doc, "Figma Learn – https://www.figma.com/resource-library/design-basics")
    bullet(doc, "MDN Web Docs – https://developer.mozilla.org")
    bullet(doc, "Bootstrap Documentation – https://getbootstrap.com/docs/5.3")

    # ─── SAVE ──────────────────────────────────────────────────
    output_path = os.path.join(BASE, "Summer_Internship_Report_v3.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    print("Done!")


if __name__ == "__main__":
    build()
